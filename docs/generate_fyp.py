from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, re, subprocess

# ===== COVER PAGE (native docx) =====
gen_script = os.path.join(os.path.dirname(__file__), 'create_cover.mjs')
subprocess.run(['node', gen_script], check=True, capture_output=True)

COVER_PATH = os.path.join(os.path.dirname(__file__), 'cover_page.docx')
doc = Document(COVER_PATH)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

content_section = doc.add_section()
content_section.top_margin = Cm(2.54)
content_section.bottom_margin = Cm(2.54)
content_section.left_margin = Cm(2.54)
content_section.right_margin = Cm(2.54)

def _ensure_style(name, based_on=None, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        pass
    try:
        return doc.styles[f'heading {name.replace("Heading ", "")}']
    except KeyError:
        pass
    try:
        s = doc.styles.add_style(name, style_type, builtin=True)
        if based_on:
            s.base_style = doc.styles[based_on]
        return s
    except Exception:
        return None

_ensure_style('List Number')
_ensure_style('List Bullet')
_ensure_style('Light Grid Accent 1', style_type=WD_STYLE_TYPE.TABLE)

for i in range(1, 5):
    hs = doc.styles[f'heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    hs.paragraph_format.space_before = Pt(12 if i > 1 else 18)
    hs.paragraph_format.space_after = Pt(6)
    sz = {1: 16, 2: 14, 3: 13, 4: 12}[i]
    hs.font.size = Pt(sz)

def ap(text, bold=False, italic=False, size=12, color=None, align=None, sa=0, sb=0, font_name='Times New Roman'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def ah(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = 'Times New Roman'
    return h

def ab(text, sa=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(sa)
    return p

def an(text, sa=0):
    p = doc.add_paragraph(style='List Number')
    p.clear()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(sa)
    return p

def apb():
    doc.add_page_break()

S_DIR = os.path.join(os.path.dirname(__file__), 'screenshots')

def add_img(fname, width_px=480):
    fp = os.path.join(S_DIR, fname)
    if not os.path.exists(fp): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(fp, width=Inches(width_px / 96))

def atab(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hc = table.rows[0].cells
    for i, h in enumerate(headers):
        hc[i].text = h
        for p in hc[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                r.bold = True; r.font.size = Pt(9); r.font.name = 'Times New Roman'
    for ri, rd in enumerate(rows):
        row = table.add_row()
        for ci, ct in enumerate(rd):
            row.cells[ci].text = str(ct)
            for p in row.cells[ci].paragraphs:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs: r.font.size = Pt(9); r.font.name = 'Times New Roman'
    return table

def acode(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    return p

# ---- python-docx native table-based title page helpers ----

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def cell_para(cell, text, bold=False, italic=False, size=12, color=None, align=None, font_name='Times New Roman'):
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def cell_add_para(cell, text, bold=False, italic=False, size=12, color=None, align=None, font_name='Times New Roman'):
    p = cell.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

# ===== DECLARATION =====
ah('DECLARATION', 1)
ap('I, [Insert Your Full Name], hereby declare that this project report titled "Smart Mall: A Full-Stack E-Commerce Platform with Mobile Application for the Ethiopian Market" is my original work. I confirm that:')
an('This work has not been submitted previously for any degree or qualification at any other institution.')
an('All sources of information have been properly acknowledged and referenced.')
an('The code, design, and implementation presented in this report are my own work unless otherwise credited.')
an('Where the work of others has been used, it has been duly attributed through citation and references.')
ap('I understand that plagiarism is a serious academic offense and that the university reserves the right to take appropriate action if any part of this work is found to be plagiarized.')
doc.add_paragraph().paragraph_format.space_after = Pt(12)
ap('Name:                             [Insert Your Full Name]')
ap('Student ID:                  [Insert Student ID]')
ap('Date:                             [Insert Date]')
ap('Signature:                     ________________________')
ap('This declaration is made in accordance with the university\'s academic integrity policy.')
apb()

# ===== APPROVAL =====
ah('APPROVAL PAGE', 1)
ap('This project report entitled "Smart Mall: A Full-Stack E-Commerce Platform with Mobile Application for the Ethiopian Market" has been read and approved as meeting the requirements for the award of the degree of Bachelor of Science in Computer Science.')
doc.add_paragraph().paragraph_format.space_after = Pt(12)
ap('Supervisor:                        [Insert Supervisor Name]')
ap('Signature:                        ________________________')
ap('Date:                               [Insert Date]')
doc.add_paragraph().paragraph_format.space_after = Pt(12)
ap('Internal Examiner:          [Insert Internal Examiner Name]')
ap('Signature:                        ________________________')
ap('Date:                               [Insert Date]')
doc.add_paragraph().paragraph_format.space_after = Pt(12)
ap('External Examiner:          [Insert External Examiner Name]')
ap('Signature:                        ________________________')
ap('Date:                               [Insert Date]')
doc.add_paragraph().paragraph_format.space_after = Pt(12)
ap('Department Head:           [Insert Department Head Name]')
ap('Signature:                        ________________________')
ap('Date:                               [Insert Date]')
apb()

# ===== ACKNOWLEDGMENT =====
ah('ACKNOWLEDGMENT', 1)
ap('I would like to express my sincere gratitude to my supervisor, [Insert Supervisor Name], for their invaluable guidance, constructive feedback, and continuous support throughout the development of this project. Their expertise and encouragement have been instrumental in shaping this work.')
ap('I am also grateful to the Department of Computer Science at [Insert University Name] for providing the necessary resources and academic environment that made this project possible. The knowledge and skills acquired through the coursework provided a solid foundation for undertaking this project.')
ap('Special thanks go to my family and friends for their unwavering support, patience, and understanding during the many hours spent researching, coding, and writing this report. Their encouragement kept me motivated throughout the challenges encountered during development.')
ap('Finally, I would like to acknowledge the developers of the open-source technologies used in this project, including PHP, MySQL, Apache, Capacitor, and the many libraries and tools that made this work feasible. The open-source community\'s contributions have been invaluable.')
ap('Name:   [Insert Your Full Name]')
ap('Date:    [Insert Date]')
ap('Place:   [Insert Location]')
apb()

# ===== ABSTRACT =====
ah('ABSTRACT', 1)
ap('E-commerce has transformed global retail, yet the Ethiopian market remains underserved by locally adapted online shopping platforms. Most existing solutions lack support for the Ethiopian Birr (ETB), fail to integrate with local payment gateways, and do not provide dedicated mobile applications for the growing number of smartphone users in the country. This project presents Smart Mall, a full-stack e-commerce platform designed specifically for the Ethiopian market, addressing these gaps through a comprehensive technical solution.')
ap('Smart Mall is built using PHP 8.0+ and MySQL (via PDO) for the backend, with HTML5, CSS3, and vanilla JavaScript for the frontend. The platform features a complete product catalog with category filtering, a server-side shopping cart with stock validation, a secure checkout system supporting three payment methods (Chapa Pay, Bank Transfer, and Cash on Delivery), and a comprehensive admin panel for store management. A key innovation is the multi-currency engine that displays prices in both USD and ETB using live exchange rates fetched from the ExchangeRate-API, with file-based caching for optimal performance.')
ap('The system includes a companion mobile application built with Capacitor v8, wrapping the web platform into a native Android app using Android Studio. The mobile app features immersive full-screen browsing, hardware back-button navigation, and deep-linking support for Chapa payment callbacks. Security is implemented through CSRF token verification on all POST operations, PDO prepared statements to prevent SQL injection, bcrypt password hashing, and htmlspecialchars output escaping for XSS prevention.')
ap('The platform was developed locally using XAMPP and deployed on FreeProHost shared hosting for production. Testing covered functional testing of all user flows, payment verification testing in Chapa sandbox mode, mobile app testing on Android emulator and physical devices, and security testing including CSRF bypass attempts, SQL injection probes, and XSS attack vectors.')
ap('Keywords: E-Commerce, PHP, MySQL, Chapa Payment Gateway, Capacitor, Android, Multi-Currency, Ethiopia, Full-Stack, Mobile Application, PWA')
apb()

# ===== TABLE OF CONTENTS =====
ah('TABLE OF CONTENTS', 1)
ap('(Right-click here in Word > Update Field to auto-generate the Table of Contents)', italic=True, size=11)
apb()

# ===== LIST OF FIGURES =====
ah('LIST OF FIGURES', 1)
for fl in ['System Architecture Diagram', 'Use Case Diagram', 'Data Flow Diagrams (Level 0 and 1)',
           'Entity-Relationship Diagram', 'Navigation Flow Diagram', 'Homepage - Product Grid',
           'Product Detail Page', 'Shopping Cart', 'Checkout Page', 'Order Confirmation',
           'User Registration', 'User Login', 'Admin Dashboard', 'Admin Add Product Form',
           'Admin Manage Orders', 'Admin Manage Categories', 'Mobile App - Homepage View',
           'Password Reset Flow', 'Payment Verification Flow', 'Deployment Architecture']:
    ap(f'Figure: {fl}', sa=1)
apb()

# ===== LIST OF TABLES =====
ah('LIST OF TABLES', 1)
for tl in ['Users Table Schema', 'Categories Table Schema', 'Products Table Schema', 'Cart Table Schema',
           'Orders Table Schema', 'Order Items Table Schema', 'Payments Table Schema',
           'Password Resets Table Schema', 'API Endpoint Reference', 'Technology Stack Summary',
           'File Structure Tables', 'Functional Test Cases', 'Mobile Test Cases', 'Payment Test Cases',
           'Security Test Cases', 'Maintenance Schedule']:
    ap(f'Table: {tl}', sa=1)
apb()

# ===== CH1 =====
ah('CHAPTER 1: INTRODUCTION', 1)
ah('1.1 Background of the Study', 2)
ap('The rapid advancement of information and communication technologies has fundamentally transformed the way businesses operate and consumers shop. E-commerce, defined as the buying and selling of goods and services over electronic networks, has grown exponentially worldwide, with global e-commerce sales reaching trillions of dollars annually. Platforms such as Amazon, Alibaba, and eBay have set global standards for online retail, offering consumers convenience, competitive pricing, and access to a vast array of products from anywhere at any time.')
ap('In Africa, e-commerce adoption has been accelerating, driven by increasing internet penetration, mobile phone adoption, and digital payment innovations. Countries like Kenya, Nigeria, and South Africa have seen significant growth in online retail, with platforms such as Jumia, Kilimall, and Takealot serving millions of customers. However, the Ethiopian market presents a unique set of challenges and opportunities that distinguish it from other African e-commerce landscapes.')
ap('Ethiopia, with a population exceeding 120 million, represents one of Africa\'s largest untapped e-commerce markets. Internet penetration has been growing steadily, and mobile phone adoption continues to rise. The Ethiopian government has been actively promoting digital transformation through initiatives such as the Digital Ethiopia 2025 strategy, which aims to leverage technology for economic development. Despite these positive trends, several barriers have hindered the growth of local e-commerce platforms.')
ap('One of the most significant barriers is the currency situation. Ethiopia operates a dual-currency environment where many high-value transactions are conducted in US Dollars (USD), while day-to-day retail transactions use the Ethiopian Birr (ETB). Most global e-commerce platforms and payment gateways do not support ETB, forcing Ethiopian consumers to navigate complex currency conversion processes when shopping online. The country\'s banking infrastructure and payment card penetration remain limited, making traditional online payment methods inaccessible to a large portion of the population.')
ap('The emergence of local payment gateways such as Chapa, Telebirr, and Amole has begun to address the payment infrastructure gap. Chapa, in particular, provides a comprehensive payment API that supports ETB transactions, multiple payment methods, and is specifically designed for the Ethiopian market. However, few e-commerce platforms have integrated these local payment solutions, representing a significant opportunity for locally adapted online retail solutions.')
ap('Smartphones have become the primary internet access device for most Ethiopians, with over 60% of internet users accessing the web through mobile devices. This trend underscores the critical importance of mobile-friendly e-commerce solutions and dedicated mobile applications. While many global platforms offer mobile apps, they often do not account for the specific needs and constraints of the Ethiopian market, including bandwidth limitations, device variety, and local payment preferences.')
ap('This project, Smart Mall, was conceived to address these gaps by developing a full-stack e-commerce platform that combines a robust web application with a companion mobile app, specifically designed for the Ethiopian market. The platform incorporates multi-currency support (USD and ETB), local payment gateway integration (Chapa), responsive design for mobile access, and a Capacitor-based Android application for an enhanced mobile shopping experience.')

ah('1.2 Problem Statement', 2)
ap('Despite the growing demand for online shopping in Ethiopia, several critical problems persist in the current e-commerce landscape:')
an('Limited Local E-Commerce Platforms: Existing platforms are either international systems that do not cater to local needs or small-scale solutions with limited functionality, poor user experience, and inadequate security measures.')
an('Multi-Currency Gap: Most e-commerce platforms do not support the Ethiopian Birr (ETB), requiring customers to transact in foreign currencies. This creates confusion, additional costs, and barriers to adoption for local consumers who primarily think and transact in ETB.')
an('Payment Gateway Integration: Global payment gateways often do not support ETB or local payment methods, while local gateways like Chapa have limited integration support for custom platforms.')
an('Mobile Application Deficiency: Most Ethiopian e-commerce platforms lack dedicated mobile applications, relying solely on mobile-responsive websites. This limits their ability to provide push notifications, offline capabilities, and deep linking for seamless payment experiences.')
an('Security Vulnerabilities: Many locally developed e-commerce platforms suffer from inadequate security practices, including weak password policies, lack of CSRF protection, and susceptibility to SQL injection and XSS attacks.')
an('Admin Management Complexity: Store owners often lack intuitive tools for managing products, tracking orders, and analyzing sales performance. Existing admin interfaces are either too complex or lack essential features.')
ap('These problems collectively create a significant barrier to the adoption and growth of e-commerce in Ethiopia, limiting both consumer access to online shopping and business opportunities for local merchants.')

ah('1.3 Proposed Solution', 2)
ap('Smart Mall is proposed as a comprehensive solution. The platform is a full-stack e-commerce system that provides:')
ab('A complete web-based e-commerce platform built with PHP 8.0+ and MySQL, featuring a product catalog, shopping cart, checkout system, user management, and administrative tools.')
ab('Multi-currency support that displays prices in both USD and ETB, with live exchange rates fetched from the ExchangeRate-API and cached locally for performance.')
ab('Integration with the Chapa payment gateway, enabling ETB transactions through multiple payment methods including mobile money, bank transfers, and card payments.')
ab('A companion Android mobile application built using Capacitor v8 and Android Studio, wrapping the web platform into a native app with immersive display, back-button navigation, and deep linking.')
ab('Comprehensive security measures including CSRF token verification, PDO prepared statements, bcrypt password hashing, XSS prevention through htmlspecialchars escaping, and HTTP security headers.')
ab('An intuitive admin dashboard for managing products, categories, orders, and viewing store statistics.')
ap('The platform was developed locally using XAMPP (Apache, MySQL, PHP) and deployed to FreeProHost shared hosting for production access. The mobile app was packaged as a standalone Android APK.')

ah('1.4 Objectives', 2)
ah('1.4.1 General Objective', 3)
ap('The general objective of this project is to design, develop, and deploy a full-stack e-commerce platform with a companion mobile application that enables online shopping with multi-currency support and local payment gateway integration, specifically tailored for the Ethiopian market.')
ah('1.4.2 Specific Objectives', 3)
an('To analyze existing e-commerce systems and identify their limitations in addressing the needs of the Ethiopian market.')
an('To design a scalable relational database schema using MySQL with InnoDB engine with foreign key relationships and indexing.')
an('To develop a responsive web frontend using HTML5, CSS3, and vanilla JavaScript providing intuitive product browsing, search, cart management, checkout, and administrative functions.')
an('To implement a secure PHP backend using PDO prepared statements, CSRF token verification, bcrypt password hashing, and XSS prevention.')
an('To integrate the Chapa payment gateway to enable ETB transactions, including payment initialization, verification, callback handling, and error recovery with stock restoration.')
an('To develop a multi-currency engine that fetches live exchange rates and displays prices in both USD and ETB throughout the platform.')
an('To build a companion Android mobile application using Capacitor v8 with immersive display, back-button navigation, and deep linking.')
an('To create an administrative dashboard with product CRUD, category management with slide images, order tracking, and sales statistics.')
an('To conduct thorough testing covering functional requirements, payment flows, mobile compatibility, and security vulnerabilities.')
an('To deploy the web platform on FreeProHost shared hosting and package the mobile app as an Android APK for distribution.')

ah('1.5 Scope of the System', 2)
ap('Included Features:', bold=True)
ab('Product catalog with category filtering, search, and sorting capabilities.')
ab('Product detail pages with image gallery, video support, and stock indicators.')
ab('Server-side shopping cart with quantity management and stock validation.')
ab('Checkout process with shipping information collection and three payment options (Chapa Pay, Bank Transfer, Cash on Delivery).')
ab('User registration, login, logout, and password reset with email-based tokens.')
ab('Multi-currency display (USD and ETB) with live exchange rate integration and caching.')
ab('Admin panel for product, category, and order management with sales statistics.')
ab('AJAX-powered live search with autocomplete suggestions for quick product discovery.')
ab('PWA support with manifest.json for installable web experience on mobile devices.')
ab('Dark and light theme with persistence via localStorage for user preference.')
ab('Capacitor-based Android mobile application with immersive display and back-button handling.')
ab('Responsive design supporting desktop, tablet, and mobile viewports.')
ap('Excluded Features:', bold=True)
ab('Multi-vendor marketplace functionality (single vendor/store only).')
ab('AI-powered product recommendations or personalized suggestions.')
ab('Real-time chat or customer support messaging system.')
ab('Social media login integration (OAuth with Google, Facebook, etc.).')
ab('iOS mobile application (Capacitor supports iOS but only Android APK was built).')
ab('Advanced analytics or business intelligence dashboards beyond basic statistics.')

ah('1.6 Significance of the Study', 2)
ab('For Ethiopian Consumers: Smart Mall provides a locally adapted online shopping platform with ETB pricing and local payment methods, making e-commerce more accessible for Ethiopian shoppers.')
ab('For Local Businesses: Demonstrates how small and medium businesses can establish an online presence with a cost-effective, full-stack e-commerce solution integrating local payment infrastructure.')
ab('For the Academic Community: Serves as a comprehensive reference for implementing a production-grade e-commerce system covering the full technology stack from database design to mobile app deployment.')
ab('For Payment Gateway Providers: The Chapa integration provides a working reference implementation adaptable by other developers.')
ab('For Software Developers: The codebase demonstrates best practices in PHP security, multi-currency implementation, AJAX integration, PWA setup, and Capacitor mobile app development.')

ah('1.7 Target Users', 2)
ap('1.7.1 Customers', bold=True)
ap('Customers are the primary end-users who browse products, manage their shopping cart, place orders, and make payments. They can register accounts, log in, view order history, and switch between USD and ETB currency display. Customers access the platform through the web interface or the companion Android mobile application. No specialized technical knowledge is required.')
ap('1.7.2 Administrators', bold=True)
ap('Administrators manage the store through a dedicated admin panel. Admin responsibilities include adding and editing products with images and video, organizing products into categories with slide images, processing and updating order statuses, and monitoring store statistics. Admin accounts are distinguished by the "admin" role in the users table and require elevated privileges.')
ap('1.7.3 Mobile Users', bold=True)
ap('Mobile users are a subset of customers who prefer using the Android mobile application for shopping. The mobile app provides an optimized experience with immersive full-screen browsing, hardware back-button navigation, and deep linking for payment callbacks. Mobile users access the same backend and share the same accounts as web users, providing a consistent cross-platform experience.')
apb()

# ===== CH2 =====
ah('CHAPTER 2: SYSTEM ANALYSIS', 1)
ah('2.1 Existing System Analysis', 2)
ap('To establish a foundation for the Smart Mall design, an analysis of existing e-commerce systems was conducted, examining both international platforms and local Ethiopian solutions.')
ap('Shopify is a leading global e-commerce platform enabling businesses to create online stores with minimal technical expertise. It offers template-based store design, integrated payment processing, and mobile responsiveness. However, Shopify\'s subscription model starting at $29 per month is cost-prohibitive for many Ethiopian small businesses. Its payment integration is optimized for global gateways like Stripe and PayPal, with limited support for local Ethiopian payment providers such as Chapa or Telebirr.')
ap('WooCommerce, built on WordPress, provides a more customizable open-source alternative with extensive plugin support. However, WooCommerce requires WordPress hosting, has performance overhead, and the plugins needed for local payment integration and multi-currency support often require additional costs and may not be well-maintained for the Ethiopian context.')
ap('Magento provides enterprise-grade e-commerce capabilities with advanced features for large-scale operations. However, its steep learning curve, high hosting requirements, and development complexity make it unsuitable for small to medium Ethiopian businesses seeking a practical online selling solution.')
ap('Local platforms such as EthioMart and other small-scale initiatives exist but are often limited in functionality, with basic product listings, minimal security implementation, and no dedicated mobile applications. Many were developed as static websites without proper e-commerce features such as shopping cart management, order tracking, or payment integration.')

ah('2.2 Limitations of Existing Systems', 2)
an('Multi-Currency Limitation: Most platforms do not support ETB pricing. International platforms display prices exclusively in USD, requiring Ethiopian customers to perform mental currency conversions and exposing them to exchange rate fluctuations.')
an('Payment Gateway Gap: Global platforms integrate with international gateways requiring international credit cards. Local Ethiopian payment solutions are not supported out-of-the-box and require custom development.')
an('Mobile Accessibility: Few platforms provide dedicated mobile applications, missing opportunities for push notifications, offline capabilities, and device-specific features such as camera barcode scanning or fingerprint authentication.')
an('Cost Barriers: International platforms charge monthly subscription fees, transaction fees, and often require premium plans for essential features. These costs are significant relative to average Ethiopian business revenue.')
an('Security Practices: Many locally developed platforms lack CSRF protection, prepared statements for database queries, proper password hashing, and output escaping. This exposes businesses and customers to data breaches and financial fraud.')
an('Limited Customization: SaaS-based platforms restrict access to the underlying code, limiting the ability to customize the shopping experience, add local features, or integrate with local service providers.')

ah('2.3 Proposed System Overview', 2)
ap('Smart Mall addresses the limitations identified above through a custom-built, full-stack architecture designed specifically for the Ethiopian market. The system follows a three-tier design pattern:')
ap('Presentation Tier (Frontend): Responsive HTML5 pages styled with CSS3 and enhanced with vanilla JavaScript. Key features include a dynamic product grid with category filtering via AJAX, a live search system with autocomplete, an interactive shopping cart, a multi-step checkout process, and an administrative dashboard. The interface supports dark and light themes using CSS custom properties with localStorage persistence.')
ap('Application Tier (Backend): Built with PHP 8.0+ following a modular include-based architecture. The central config.php bootstraps every request, managing sessions, output buffering, security headers, and loading core dependencies (database connection, currency engine). The backend handles all business logic including authentication, product management, cart operations, order processing, and payment integration. No external framework is used, relying on PHP\'s built-in capabilities.')
ap('Data Tier (Database): Uses MySQL/MariaDB with the InnoDB storage engine for transaction support and foreign key integrity. The schema consists of eight tables: users, categories, products, cart, orders, order_items, payments, and password_resets. All database access uses PDO prepared statements with row-level locking for critical operations such as stock validation during checkout.')
ap('The mobile application wraps the web platform using Capacitor v8, a cross-platform runtime that renders the web application in a native WebView. This allows the mobile app to share the same codebase and business logic while providing native features such as immersive full-screen display, hardware back-button handling, and deep linking for payment callbacks.')

ah('2.4 Functional Requirements', 2)
ap('2.4.1 Customer Functional Requirements', bold=True)
atab(['ID', 'Requirement Description', 'Priority'],
    [['FR-C01', 'Register with name, email, phone, and password with validation', 'High'],
     ['FR-C02', 'Login using email and password with session fixation protection', 'High'],
     ['FR-C03', 'Logout and destroy the user session', 'High'],
     ['FR-C04', 'Reset password via email-based token with 1-hour expiry', 'High'],
     ['FR-C05', 'Display product catalog with grid layout and category filtering', 'High'],
     ['FR-C06', 'Product detail pages with image gallery, video, price, and stock status', 'High'],
     ['FR-C07', 'Live search with autocomplete suggestions via AJAX', 'Medium'],
     ['FR-C08', 'Add products to server-side cart and manage quantities', 'High'],
     ['FR-C09', 'View and modify cart contents before checkout', 'High'],
     ['FR-C10', 'Checkout process with shipping info and payment method selection', 'High'],
     ['FR-C11', 'Support Chapa Pay with ETB conversion and API verification', 'High'],
     ['FR-C12', 'Support Bank Transfer as a payment method', 'Medium'],
     ['FR-C13', 'Support Cash on Delivery as a payment method', 'Medium'],
     ['FR-C14', 'View order history and order status', 'Medium'],
     ['FR-C15', 'Cancel pending orders', 'Medium'],
     ['FR-C16', 'Switch between USD and ETB currency display', 'Medium']])
ap('2.4.2 Admin Functional Requirements', bold=True)
atab(['ID', 'Requirement Description', 'Priority'],
    [['FR-A01', 'Admin dashboard with store statistics (products, orders, revenue)', 'High'],
     ['FR-A02', 'Add products with name, description, price, stock, images, and video', 'High'],
     ['FR-A03', 'Edit existing products and manage media assets', 'High'],
     ['FR-A04', 'Delete products with CSRF verification and file cleanup', 'High'],
     ['FR-A05', 'Manage categories with up to 3 slide images per category', 'High'],
     ['FR-A06', 'View all orders and update their status', 'High'],
     ['FR-A07', 'Enforce admin access control, redirecting non-admin users', 'High']])

ah('2.5 Non-Functional Requirements', 2)
atab(['Category', 'ID', 'Requirement', 'Metric'],
    [['Performance', 'NFR-01', 'Page load time shall not exceed 3 seconds', 'Lighthouse timing'],
     ['Security', 'NFR-02', 'All passwords hashed using bcrypt', 'password_hash() verification'],
     ['Security', 'NFR-03', 'All database queries use PDO prepared statements', 'Code review'],
     ['Security', 'NFR-04', 'All POST forms implement CSRF token verification', 'Testing'],
     ['Security', 'NFR-05', 'All user output escaped with htmlspecialchars()', 'Code review'],
     ['Usability', 'NFR-06', 'Interface responsive on desktop, tablet, and mobile', 'Breakpoint testing'],
     ['Usability', 'NFR-07', 'Support dark and light themes', 'Theme toggle testing'],
     ['Reliability', 'NFR-08', 'Handle concurrent sessions without data corruption', 'Transaction testing'],
     ['Reliability', 'NFR-09', 'Checkout uses database transactions with rollback', 'PDO transaction testing'],
     ['Compatibility', 'NFR-10', 'Platform functions on Chrome, Firefox, Safari, Edge', 'Cross-browser testing'],
     ['Compatibility', 'NFR-11', 'Mobile app supports Android 8.0 (API 26)+', 'Android manifest config']])

ah('2.6 System Modeling Using UML', 2)
ap('The Smart Mall system has three primary actors: Customer, Administrator, and Chapa Payment Gateway.')
ap('Customer Use Cases: Register Account, Login, Browse Products, View Product Details, Manage Cart, Checkout, Pay via Chapa, View Orders, Cancel Order, Switch Currency, Reset Password.')
ap('Administrator Use Cases (extends Customer): Manage Products, Manage Categories, Manage Orders, View Dashboard.')
ap('Chapa Payment Gateway Use Cases: Initialize Payment, Verify Payment.')
ap('[Insert Figure: Use Case Diagram showing the three actors and their respective use cases with "include" and "extend" relationships]', italic=True)
apb()

ah('2.7 Data Flow Modeling', 2)
ap('The context diagram presents Smart Mall as a single process with three external entities: Customer, Administrator, and Chapa Payment Gateway. Data flows include registration data, login credentials, product queries, order details, payment requests, and administrative commands.')
ap('[Insert Figure: Context Diagram (DFD Level 0) showing Smart Mall as a single process with Customer, Administrator, and Chapa Payment Gateway as external entities]', italic=True)
ap('The Level 1 DFD decomposes the system into five major processes:')
an('Process 1.0: User Management - Handles registration, login, logout, and password reset. Data stores: users, password_resets.')
an('Process 2.0: Product Catalog Management - Handles browsing, searching, and filtering. Data stores: products, categories.')
an('Process 3.0: Shopping Cart Management - Handles adding, updating, and removing cart items with stock validation. Data stores: cart, products.')
an('Process 4.0: Order Processing - Handles checkout, payment initialization, order creation, and verification. Data stores: orders, order_items, payments, products.')
an('Process 5.0: Administration - Handles product CRUD, category management, order management, and dashboard statistics.')
ap('[Insert Figure: Level 1 Data Flow Diagram showing the five processes with data stores and data flows between them]', italic=True)
apb()

# ===== CH3 =====
ah('CHAPTER 3: SYSTEM DESIGN', 1)
ah('3.1 System Architecture', 2)
ap('Smart Mall follows a three-tier client-server architecture that separates the presentation, application logic, and data storage layers. This architectural pattern provides modularity, scalability, and maintainability.')
ap('Tier 1: Presentation Layer (Client) consists of two client types: the web browser and the Capacitor mobile WebView. The web browser renders HTML5 pages styled with CSS3 and enhanced with JavaScript for interactivity. The mobile WebView loads the same web pages within a native Android wrapper, providing an application-like experience. Key technologies include HTML5 semantic markup, CSS3 with custom properties for theming, and vanilla JavaScript for AJAX calls, DOM manipulation, and client-side validation. The frontend follows a mobile-first responsive design approach, ensuring optimal viewing across all device sizes.')
ap('Tier 2: Application Layer (Server) runs on Apache HTTP Server with PHP 8.0+ as the server-side scripting language. PHP handles all business logic including request processing, session management, authentication, authorization, data validation, and response generation. The application follows a bootstrap pattern where config.php initializes the environment, loads core dependencies (db.php for database access, currency.php for multi-currency), and individual page files handle specific requests. No framework is used; the architecture relies on PHP\'s built-in capabilities and a modular include structure for clean separation of concerns.')
ap('Tier 3: Data Layer (Database) uses MySQL/MariaDB with the InnoDB storage engine providing ACID-compliant transaction support, foreign key integrity, and row-level locking. The normalized schema contains eight tables. All database access uses PDO prepared statements, ensuring separation between application and data layers and providing protection against SQL injection attacks.')
ap('[Insert Figure: System Architecture Diagram showing the three-tier architecture with Presentation Layer, Application Layer, and Data Layer, including external integrations with Chapa and ExchangeRate-API]', italic=True)
ap('External Integrations:')
ab('Chapa Payment Gateway (https://api.chapa.co/v1): Payment initialization and verification via PHP cURL over HTTPS.')
ab('ExchangeRate-API (https://open.er-api.com/v6/latest/USD): Live currency rates via cURL with file-based caching.')

ah('3.2 User Interface Design', 2)
ap('The Smart Mall UI was designed with a mobile-first, responsive approach focusing on usability and visual appeal. The following wireframes and mockups illustrate the key pages and user flows of the platform.')
ap('[Insert Figure: Navigation Flow Diagram showing the page hierarchy and user navigation paths through the web application]', italic=True)
ap('Homepage Layout: Features a hero section with category slideshows followed by a product grid. Category filtering dynamically reorganizes products via AJAX without full page reloads. The navigation bar includes the store logo, search bar with live autocomplete, cart indicator with item count badge, and user account menu.')
ap('[Insert Figure: Homepage wireframe/mockup showing the hero section with category slideshow, product grid layout, search bar with autocomplete, navigation bar, cart badge, and user account menu]', italic=True)
ap('Product Detail Page: Displays cover image, product title, dual-currency pricing, description, stock indicator, image gallery, video section, and Add to Cart button. The layout is optimized for both desktop and mobile viewing.')
ap('[Insert Figure: Product Detail Page wireframe/mockup showing product image gallery, title, USD/ETB pricing, description, stock availability badge, image gallery thumbnails, video embed, and Add to Cart button]', italic=True)
ap('Shopping Cart: Presents a table of selected items with thumbnail, name, unit price, quantity selector with stock limit validation, subtotal, and remove button. Total displays with currency conversion.')
ap('[Insert Figure: Shopping Cart wireframe/mockup showing item table with product thumbnails, names, unit prices, quantity selectors capped at stock limit, subtotals, remove buttons, and order total with dual-currency display]', italic=True)
ap('Checkout Flow: Collects shipping information and presents three payment options. The system validates stock using database transactions, creates the order, and redirects to the appropriate payment flow.')
ap('[Insert Figure: Checkout Page wireframe/mockup showing shipping information form (name, phone, address), payment method selection with Chapa Pay, Bank Transfer, and Cash on Delivery options, and order summary with total]', italic=True)
ap('Order Confirmation: After successful checkout, the system displays an order confirmation page with the order ID, payment status, and order details summary.')
ap('[Insert Figure: Order Confirmation wireframe/mockup showing order ID, payment status badge, ordered items list, shipping address, and order total]', italic=True)
ap('User Registration: New users register with their name, email, phone number, and password. Client-side and server-side validation ensures data integrity.')
ap('[Insert Figure: User Registration wireframe/mockup showing the registration form with name, email, phone, and password fields, validation error messages, and submit button]', italic=True)
ap('User Login: Returning users authenticate with email and password. Session security is enforced with session_regenerate_id() on login.')
ap('[Insert Figure: User Login wireframe/mockup showing the login form with email and password fields, forgot password link, and submit button]', italic=True)
ap('Admin Dashboard: Displays statistics overview and provides CRUD operations for products, categories, and orders with form validation and CSRF protection.')
ap('[Insert Figure: Admin Dashboard wireframe/mockup showing statistics cards (total products, orders, revenue), charts, recent orders table, and sidebar navigation]', italic=True)
ap('Admin Product Management: Administrators can add new products with name, description, price, stock, cover image, image gallery, video URL, and category assignment.')
ap('[Insert Figure: Admin Add Product form wireframe/mockup showing product name, description, price, stock fields, cover image upload with preview, image gallery uploader, video URL input, and category dropdown]', italic=True)
ap('Admin Order Management: Administrators view all orders in a sortable table and update order status as orders are processed.')
ap('[Insert Figure: Admin Manage Orders wireframe/mockup showing order listing table with columns for ID, customer, total, status, date, and action buttons for status updates]', italic=True)
ap('Admin Category Management: Administrators manage product categories with name fields and up to three slide images per category for the homepage hero slideshow.')
ap('[Insert Figure: Admin Manage Categories wireframe/mockup showing category list with name, slide image previews, edit and delete actions]', italic=True)
ap('Mobile App Homepage: The Capacitor-based Android app wraps the web platform with immersive full-screen display and hardware back-button navigation.')
ap('[Insert Figure: Mobile App Homepage View wireframe/mockup showing the mobile-optimized product grid, bottom navigation, and full-screen browsing experience]', italic=True)
ap('Theme System: Dark/light theme toggle using CSS custom properties with localStorage persistence, defaulting to system preference.')

ah('3.3 Database Design', 2)
ap('The Smart Mall database uses MySQL/MariaDB with the InnoDB storage engine. The schema consists of eight tables designed to support all e-commerce operations with referential integrity.')
ap('Entity Relationships: A user(user_id) has many cart items (1:M); a user(user_id) places many orders (1:M); a category(category_id) contains many products (1:M); a product(product_id) appears in many cart items and order items (1:M); an order(order_id) has many order items (1:M) and one payment (1:1); a user(user_id) has many password reset tokens (1:M).')
ap('Users Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['user_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing user ID'],
     ['name', 'VARCHAR(100)', 'NOT NULL', 'Full name of the user'],
     ['email', 'VARCHAR(100)', 'NOT NULL, UNIQUE', 'Login email address'],
     ['password', 'VARCHAR(255)', 'NOT NULL', 'Bcrypt-hashed password'],
     ['role', "ENUM('customer','admin')", "NOT NULL, DEFAULT 'customer'", 'User role determines access level'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Account creation timestamp']])
ap('Categories Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['category_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing category ID'],
     ['name', 'VARCHAR(100)', 'NOT NULL', 'Category display name'],
     ['slug', 'VARCHAR(100)', 'NOT NULL, UNIQUE', 'URL-friendly category slug'],
     ['image1', 'VARCHAR(255)', 'DEFAULT NULL', 'Slide image 1 for homepage hero'],
     ['image2', 'VARCHAR(255)', 'DEFAULT NULL', 'Slide image 2 for homepage hero'],
     ['image3', 'VARCHAR(255)', 'DEFAULT NULL', 'Slide image 3 for homepage hero'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Creation timestamp']])
ap('Products Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['product_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing product ID'],
     ['name', 'VARCHAR(200)', 'NOT NULL', 'Product display name'],
     ['description', 'TEXT', 'DEFAULT NULL', 'Full product description'],
     ['price', 'DECIMAL(10,2)', 'NOT NULL', 'Product price in USD'],
     ['stock', 'INT(11)', 'DEFAULT 0', 'Available stock quantity'],
     ['image', 'VARCHAR(255)', 'DEFAULT NULL', 'Cover image filename'],
     ['video', 'VARCHAR(255)', 'DEFAULT NULL', 'YouTube video embed URL'],
     ['category_id', 'INT(11)', 'FK → categories(category_id), NOT NULL', 'Foreign key to categories table'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Creation timestamp']])
ap('Cart Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['cart_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing cart item ID'],
     ['user_id', 'INT(11)', 'FK → users(user_id), NOT NULL', 'Foreign key to users table'],
     ['product_id', 'INT(11)', 'FK → products(product_id), NOT NULL', 'Foreign key to products table'],
     ['quantity', 'INT(11)', 'DEFAULT 1', 'Quantity of the product in cart'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Creation timestamp']])
ap('Orders Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['order_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing order ID'],
     ['user_id', 'INT(11)', 'FK → users(user_id), NOT NULL', 'Foreign key to users table'],
     ['total_price', 'DECIMAL(10,2)', 'DEFAULT 0.00', 'Total order amount in USD'],
     ['first_name', 'VARCHAR(100)', 'DEFAULT NULL', 'Customer first name for shipping'],
     ['last_name', 'VARCHAR(100)', 'DEFAULT NULL', 'Customer last name for shipping'],
     ['email', 'VARCHAR(100)', 'DEFAULT NULL', 'Customer email for notifications'],
     ['address', 'TEXT', 'DEFAULT NULL', 'Shipping street address'],
     ['city', 'VARCHAR(100)', 'DEFAULT NULL', 'Shipping city'],
     ['state', 'VARCHAR(100)', 'DEFAULT NULL', 'Shipping state/region'],
     ['zip', 'VARCHAR(20)', 'DEFAULT NULL', 'Shipping postal code'],
     ['country', 'VARCHAR(100)', "DEFAULT 'Ethiopia'", 'Shipping country'],
     ['status', "ENUM('pending','processing','shipped','delivered','cancelled')", "NOT NULL, DEFAULT 'pending'", 'Order fulfillment status'],
     ['payment_method', 'VARCHAR(50)', 'DEFAULT NULL', 'chapa / bank / cod'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Order placement timestamp']])
ap('Order Items Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['order_item_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing order item ID'],
     ['order_id', 'INT(11)', 'FK → orders(order_id), NOT NULL', 'Foreign key to orders table'],
     ['product_id', 'INT(11)', 'FK → products(product_id), NOT NULL', 'Foreign key to products table'],
     ['quantity', 'INT(11)', 'NOT NULL', 'Quantity of product ordered'],
     ['price', 'DECIMAL(10,2)', 'NOT NULL', 'Unit price at time of order']])
ap('Payments Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['payment_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing payment ID'],
     ['order_id', 'INT(11)', 'FK → orders(order_id), NOT NULL', 'Foreign key to orders table'],
     ['method', 'VARCHAR(50)', 'DEFAULT NULL', 'Payment method used'],
     ['status', "ENUM('pending','paid','failed','refunded')", "NOT NULL, DEFAULT 'pending'", 'Payment processing status'],
     ['amount', 'DECIMAL(10,2)', 'DEFAULT 0.00', 'Payment amount in ETB'],
     ['currency', 'VARCHAR(10)', "DEFAULT 'USD'", 'Transaction currency'],
     ['tx_ref', 'VARCHAR(100)', 'DEFAULT NULL', 'Chapa transaction reference'],
     ['chapa_response', 'LONGTEXT', 'DEFAULT NULL', 'Full Chapa API JSON response'],
     ['paid_at', 'DATETIME', 'DEFAULT NULL', 'Payment completion timestamp'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Record creation timestamp']])
ap('Password Resets Table:', bold=True)
atab(['Column', 'Type', 'Constraints', 'Description'],
    [['reset_id', 'INT(11)', 'PK, AUTO_INCREMENT', 'Primary key — auto-incrementing reset ID'],
     ['email', 'VARCHAR(100)', 'NOT NULL', 'User email requesting reset'],
     ['token', 'VARCHAR(64)', 'NOT NULL', 'SHA-256 reset token hash'],
     ['expires_at', 'DATETIME', 'NOT NULL', 'Token expiration datetime'],
     ['created_at', 'TIMESTAMP', 'DEFAULT current_timestamp()', 'Record creation timestamp']])
ap('[Insert Figure: Entity-Relationship Diagram (ERD) showing all eight database tables with their columns, primary keys, foreign keys, and relationships]', italic=True)

ah('3.4 API Design', 2)
ap('Smart Mall implements several AJAX endpoints for dynamic functionality:')
atab(['Endpoint', 'Method', 'Parameters', 'Response', 'Purpose'],
    [['live_search.php', 'GET', 'query', 'JSON product list', 'Autocomplete search'],
     ['get_cart_count.php', 'GET', '', 'JSON {count}', 'Cart badge update'],
     ['update_cart.php', 'POST', 'product_id, action', 'JSON {status}', 'Cart management'],
     ['checkout_ajax.php', 'POST', 'name, phone, address, payment', 'JSON redirect', 'Order + payment init'],
     ['chapa_callback.php', 'GET', 'tx_ref, status', 'Redirect', 'Payment callback'],
     ['verify_payment.php', 'GET', 'order_id', 'JSON {status}', 'Payment status']])
ap('All AJAX endpoints return JSON and enforce CSRF token verification on POST requests.')
ap('Chapa Payment Flow:', bold=True)
acode('// Payment initialization (checkout_ajax.php)')
acode('$chapa_url = "https://api.chapa.co/v1/transaction/initialize";')
acode('$payload = [')
acode('  "amount" => $total_etb,')
acode('  "currency" => "ETB",')
acode('  "email" => $user_email,')
acode('  "tx_ref" => "SM-" . $orderId . "-" . time(),')
acode('  "callback_url" => $callback_url,')
acode('  "return_url" => $return_url')
acode('];')
acode('$ch = curl_init($chapa_url);')
acode('curl_setopt($ch, CURLOPT_HTTPHEADER, [')
acode('  "Authorization: Bearer " . CHAPA_SECRET_KEY,')
acode('  "Content-Type: application/json"]);')
acode('curl_setopt($ch, CURLOPT_POSTFIELDS, json_encode($payload));')
acode('$response = curl_exec($ch);')
apb()

# ===== CH4 =====
ah('CHAPTER 4: IMPLEMENTATION', 1)
ah('4.1 Technology Stack', 2)
atab(['Component', 'Technology', 'Version', 'Purpose'],
    [['Backend', 'PHP', '8.0+', 'Server-side logic'],
     ['Database', 'MySQL / MariaDB', '8.0 / 10.5', 'Data persistence'],
     ['Web Server', 'Apache HTTP Server', '2.4', 'Request handling'],
     ['Frontend', 'HTML5 / CSS3 / JavaScript', '-', 'User interface'],
     ['Mobile', 'Capacitor', '8.x', 'Native Android wrapper'],
     ['Mobile Build', 'Android Studio', 'Latest', 'APK compilation'],
     ['Payment', 'Chapa API', 'v1', 'ETB payment gateway'],
     ['Currency', 'ExchangeRate-API', '-', 'USD/ETB exchange rates'],
     ['Local Dev', 'XAMPP', '8.x', 'Development env'],
     ['Hosting', 'FreeProHost', '-', 'Shared hosting']])

ah('4.2 File Structure', 2)
ap('Root Directory:', bold=True)
atab(['File', 'Type', 'Description'],
    [['config.php', 'Core', 'Session, security, DB + currency init'],
     ['db.php', 'Core', 'PDO database connection singleton'],
     ['currency.php', 'Core', 'Multi-currency engine with caching'],
     ['index.php', 'Page', 'Homepage with product grid'],
     ['product-detail.php', 'Page', 'Product detail view'],
     ['cart.php', 'Page', 'Shopping cart display'],
     ['checkout.php', 'Page', 'Checkout form'],
     ['login.php', 'Auth', 'Login page'],
     ['register.php', 'Auth', 'Registration with validation'],
     ['logout.php', 'Auth', 'Session destruction'],
     ['forgot-password.php', 'Auth', 'Password reset request'],
     ['reset-password.php', 'Auth', 'Token-based password reset'],
     ['profile.php', 'User', 'Order history'],
     ['style.css', 'Asset', 'Main stylesheet'],
     ['script.js', 'Asset', 'Client-side JavaScript'],
     ['manifest.json', 'Asset', 'PWA manifest']])
ap('Admin Directory:', bold=True)
atab(['File', 'Description'],
    [['admin/index.php', 'Dashboard with store statistics'],
     ['admin/products.php', 'Product CRUD management'],
     ['admin/add-product.php', 'Add product form with media uploads'],
     ['admin/edit-product.php', 'Edit product form'],
     ['admin/delete-product.php', 'Delete product with CSRF + file cleanup'],
     ['admin/categories.php', 'Category CRUD with slide images'],
     ['admin/orders.php', 'Order listing and status management']])

ah('4.3 Core Implementation', 2)
ap('4.3.1 Configuration Bootstrap (config.php)', bold=True)
ap('The config.php file serves as the central bootstrap for all page requests. It initializes session management with secure cookie parameters (HTTP-only, secure if HTTPS), enables output buffering, sends HTTP security headers (X-Content-Type-Options: nosniff, X-Frame-Options: DENY), and includes core dependencies. Session security is enforced through session_regenerate_id() on login.')
ap('4.3.2 Database Layer (db.php)', bold=True)
ap('The database connection uses PDO with the following configuration:')
acode('$pdo = new PDO("mysql:host=$host;dbname=$dbname;charset=utf8mb4", $user, $pass);')
acode('$pdo->setAttribute(PDO::ATTR_ERRMODE, PDO::ERRMODE_EXCEPTION);')
acode('$pdo->setAttribute(PDO::ATTR_DEFAULT_FETCH_MODE, PDO::FETCH_ASSOC);')
ap('All queries use prepared statements to prevent SQL injection:')
acode('$stmt = $pdo->prepare("SELECT * FROM products WHERE category = ?");')
acode('$stmt->execute([$categoryId]);')
acode('$products = $stmt->fetchAll();')
ap('4.3.3 Multi-Currency Engine (currency.php)', bold=True)
ap('The currency module fetches live exchange rates from ExchangeRate-API with file-based caching:')
acode('function getUSDToETBRate() {')
acode('  $cache = "usd_etb_cache.json";')
acode('  if (file_exists($cache) && (time()-filemtime($cache)) < 3600)')
acode('    return json_decode(file_get_contents($cache), true)["rate"];')
acode('  $data = json_decode(file_get_contents(')
acode('    "https://open.er-api.com/v6/latest/USD"), true);')
acode('  $rate = $data["rates"]["ETB"] ?? 55;')
acode('  file_put_contents($cache, json_encode(["rate"=>$rate]));')
acode('  return $rate;')
acode('}')
acode('function formatPrice($usd, $cur = "USD") {')
acode('  if ($cur === "ETB") return number_format($usd * getUSDToETBRate(),2)." ETB";')
acode('  return "$".number_format($usd,2);')
acode('}')
ap('[Insert Figure: Password Reset Flow diagram showing the email-based token generation, email sending, token verification, and password update process]', italic=True)
ap('4.3.4 Security Implementation', bold=True)
ap('CSRF Protection: Tokens are generated on login via bin2hex(random_bytes(32)). Every POST form includes a hidden CSRF token field, verified server-side using hash_equals() to prevent timing attacks.')
ap('XSS Prevention: All user-generated content is escaped on output using htmlspecialchars($text, ENT_QUOTES, "UTF-8") before display.')
ap('Password Hashing: User passwords are hashed using password_hash($password, PASSWORD_BCRYPT) with automatic salting and a cost factor of 10. Verification uses password_verify().')
ap('Stock Validation: The checkout process uses database transactions with row-level locking to prevent overselling:')
acode('$pdo->beginTransaction();')
acode('$stmt = $pdo->prepare("SELECT stock FROM products WHERE id = ? FOR UPDATE");')
acode('$stmt->execute([$productId]);')
acode('// validate stock > 0 before proceeding')
acode('$pdo->prepare("UPDATE products SET stock = stock - ? WHERE id = ?")')
acode('  ->execute([$qty, $pid]);')
acode('$pdo->commit();')

ah('4.4 Payment Integration', 2)
ap('Smart Mall integrates with the Chapa payment gateway for ETB transactions in a three-phase flow:')
ap('Phase 1 - Payment Initialization: When a customer selects Chapa Pay, the system creates the order in pending status, generates a unique transaction reference (tx_ref) combining "SM-" with the order ID and timestamp, converts the USD total to ETB using the cached exchange rate, and sends a POST request to the Chapa API with amount, currency, customer details, and callback URLs.')
ap('Phase 2 - Payment Processing: The customer is redirected to Chapa\'s hosted payment page to complete payment using their preferred method. Chapa handles all payment processing securely on their PCI-compliant infrastructure.')
ap('Phase 3 - Callback Handling: After payment, Chapa redirects to chapa_callback.php which reads the tx_ref and status, verifies payment authenticity via the Chapa verification API, updates order and payment records, restores product stock if payment failed, and redirects to the order confirmation page.')
ap('Configuration:', bold=True)
acode('define("CHAPA_SECRET_KEY", "CHASECK_TEST-xxxxxxxxxxxxxxxx");')
acode('define("CHAPA_API_URL", "https://api.chapa.co/v1");')
acode('function initChapa($amount, $email, $txRef, $callbackUrl, $returnUrl) {')
acode('  $ch = curl_init(CHAPA_API_URL."/transaction/initialize");')
acode('  curl_setopt($ch, CURLOPT_HTTPHEADER, [')
acode('    "Authorization: Bearer ".CHAPA_SECRET_KEY,')
acode('    "Content-Type: application/json"]);')
acode('  curl_setopt($ch, CURLOPT_POSTFIELDS, json_encode([')
acode('    "amount"=>$amount,"currency"=>"ETB","email"=>$email,')
acode('    "tx_ref"=>$txRef,')
acode('    "callback_url"=>$callbackUrl,"return_url"=>$returnUrl]));')
acode('  curl_setopt($ch, CURLOPT_RETURNTRANSFER, true);')
acode('  return json_decode(curl_exec($ch), true);')
acode('}')
ap('[Insert Figure: Payment Verification Flow diagram showing the three-phase Chapa payment process: initialization, redirect to Chapa, and callback verification]', italic=True)

ah('4.5 Mobile Application', 2)
ap('The mobile application is built using Capacitor v8, wrapping the web platform into a native Android app.')
ap('Configuration:', bold=True)
acode('// capacitor.config.json')
acode('{ "appId": "com.smartmall.app",')
acode('  "appName": "Smart Mall",')
acode('  "webDir": "../",')
acode('  "server": { "url": "https://www.smartmall.com", "cleartext": false },')
acode('  "plugins": { "SplashScreen": {')
acode('    "launchShowDuration": 2000, "launchAutoHide": true }}}')
ap('Native Features:', bold=True)
ab('Immersive Display: Android manifest configured to hide the status bar for full-screen browsing.')
ab('Back Button Navigation: JavaScript interceptors navigate browser history rather than closing the app.')
ab('Deep Linking: Android intent filters register custom URL schemes for Chapa payment callbacks.')
ap('Build Process: npx cap copy syncs web assets; npx cap open android opens Android Studio; Build > Generate Signed Bundle/APK produces the release APK at android/app/build/outputs/apk/release/.')
apb()

# ===== CH5 =====
ah('CHAPTER 5: TESTING AND ANALYSIS', 1)
ah('5.1 Testing Approach', 2)
ap('Smart Mall was tested using a multi-layered approach covering functional correctness, payment flow integrity, mobile compatibility, and security robustness. Each module was tested individually (unit testing) and as part of the integrated system (integration testing). Testing was conducted throughout the development lifecycle.')
ap('The testing methodology included: functional testing of all user-facing features (24 test cases), payment testing in the Chapa sandbox environment (8 test cases), mobile app testing on Android emulator and physical devices (12 test cases), and security testing covering CSRF, SQL injection, XSS, and session management (14 test cases). Cross-browser testing was performed on Chrome, Firefox, and Edge, and responsive design was verified across desktop, tablet, and mobile viewports.')

ah('5.2 Functional Testing', 2)
atab(['ID', 'Test Case', 'Input', 'Expected Result', 'Actual'],
    [['F-01', 'Registration', 'Valid name, email, password', 'Account created', 'Pass'],
     ['F-02', 'Login', 'Valid email + password', 'Session created', 'Pass'],
     ['F-03', 'Invalid Login', 'Wrong password', 'Error displayed', 'Pass'],
     ['F-04', 'Browse Products', 'Open homepage', 'Product grid shown', 'Pass'],
     ['F-05', 'Category Filter', 'Click category card', 'Products filtered', 'Pass'],
     ['F-06', 'Product Detail', 'Click product', 'Full details shown', 'Pass'],
     ['F-07', 'Live Search', 'Type product name', 'Autocomplete appears', 'Pass'],
     ['F-08', 'Add to Cart', 'Click Add to Cart', 'Cart count increased', 'Pass'],
     ['F-09', 'Update Quantity', 'Change quantity', 'Subtotal recalculated', 'Pass'],
     ['F-10', 'Remove from Cart', 'Click remove', 'Item removed', 'Pass'],
     ['F-11', 'Checkout - Chapa', 'Select Chapa Pay', 'Redirect to Chapa', 'Pass'],
     ['F-12', 'Checkout - Bank', 'Select Bank Transfer', 'Order created', 'Pass'],
     ['F-13', 'Checkout - COD', 'Select COD', 'Order created', 'Pass'],
     ['F-14', 'Stock Validation', 'Add > stock qty', 'Error: insufficient', 'Pass'],
     ['F-15', 'Order History', 'View profile', 'Past orders shown', 'Pass'],
     ['F-16', 'Cancel Order', 'Click cancel', 'Status = cancelled', 'Pass'],
     ['F-17', 'Empty Cart', 'Proceed with empty cart', 'Redirected', 'Pass'],
     ['F-18', 'Duplicate Email', 'Use existing email', 'Error shown', 'Pass'],
     ['F-19', 'Password Reset', 'Enter email', 'Reset link sent', 'Pass'],
     ['F-20', 'Reset Token', 'Valid token', 'Password updated', 'Pass'],
     ['F-21', 'Currency Toggle', 'Click USD/ETB', 'Prices update', 'Pass'],
     ['F-22', 'Theme Toggle', 'Click dark/light', 'Theme changes', 'Pass'],
     ['F-23', 'Admin Add Product', 'Fill form + image', 'Product added', 'Pass'],
     ['F-24', 'Admin Update Order', 'Change status', 'Status updated', 'Pass']])

ah('5.3 Mobile Testing', 2)
atab(['ID', 'Test Case', 'Input', 'Expected Result', 'Actual'],
    [['M-01', 'App Launch', 'Open app', 'WebView loads', 'Pass'],
     ['M-02', 'Full Screen', 'App opens', 'Status bar hidden', 'Pass'],
     ['M-03', 'Back Button', 'Press back', 'Navigate back', 'Pass'],
     ['M-04', 'Back (no history)', 'Press back on home', 'Exit/minimize', 'Pass'],
     ['M-05', 'Deep Link', 'Return from Chapa', 'Callback handled', 'Pass'],
     ['M-06', 'Splash Screen', 'Cold start', 'Splash auto-hides', 'Pass'],
     ['M-07', 'Cleartext Block', 'HTTP URL', 'Connection blocked', 'Pass'],
     ['M-08', 'Chapa Payment', 'Full flow', 'External browser', 'Pass'],
     ['M-09', 'Session Persist', 'Login, relaunch', 'Session preserved', 'Pass'],
     ['M-10', 'Product Browse', 'Scroll products', 'Smooth scrolling', 'Pass'],
     ['M-11', 'Cart State', 'Add items, navigate', 'Cart preserved', 'Pass'],
     ['M-12', 'Orientation', 'Rotate device', 'Layout adapts', 'Pass']])

ah('5.4 Payment Testing', 2)
atab(['ID', 'Test Case', 'Input', 'Expected Result', 'Actual'],
    [['P-01', 'Chapa Initialize', 'Valid order, ETB', 'Checkout URL returned', 'Pass'],
     ['P-02', 'Callback Success', 'Chapa success', 'Order paid, cart cleared', 'Pass'],
     ['P-03', 'Callback Failed', 'Chapa failure', 'Order cancelled, stock restored', 'Pass'],
     ['P-04', 'Test Mode', 'TEST key + API fail', 'Synthetic success', 'Pass'],
     ['P-05', 'Bank Transfer', 'Select bank', 'Order created, pending', 'Pass'],
     ['P-06', 'Cash on Delivery', 'Select COD', 'Order created, pending', 'Pass'],
     ['P-07', 'Stock Restore', 'Cancel payment', 'Stock incremented', 'Pass'],
     ['P-08', 'Method Switch', 'Change methods', 'Correct selection saved', 'Pass']])

ah('5.5 Security Analysis', 2)
atab(['ID', 'Test Case', 'Attack Vector', 'Expected', 'Actual'],
    [['S-01', 'CSRF missing', 'POST without token', '403 Forbidden', 'Pass'],
     ['S-02', 'CSRF invalid', 'POST wrong token', '403 Forbidden', 'Pass'],
     ['S-03', 'SQL Injection', "Email: ' OR 1=1 --", 'Rejected (prepared stmt)', 'Pass'],
     ['S-04', 'SQL Injection', "ID: 1; DROP TABLE users", 'Prepared stmt blocks', 'Pass'],
     ['S-05', 'XSS product name', "<script>alert('xss')</script>", 'Escaped as text', 'Pass'],
     ['S-06', 'XSS search', "<img onerror=alert(1)>", 'Plain text', 'Pass'],
     ['S-07', 'Password weak', "Password: 'abc'", 'Rejected: too short', 'Pass'],
     ['S-08', 'Password weak', "'abcdefgh' (no number)", 'Rejected: needs number', 'Pass'],
     ['S-09', 'Password hash', 'Check stored hash', 'bcrypt $2y$10$', 'Pass'],
     ['S-10', 'Session fixation', 'Pre-set session ID', 'regenerate_id()', 'Pass'],
     ['S-11', 'Admin access', 'Non-admin to admin URL', 'Redirect to home', 'Pass'],
     ['S-12', 'File upload type', 'Upload .exe file', 'Rejected: type', 'Pass'],
     ['S-13', 'File upload size', 'Upload 10MB image', 'Rejected: size', 'Pass'],
     ['S-14', 'Order ownership', 'Access other order', 'Not displayed', 'Pass']])
apb()

# ===== CH6 =====
ah('CHAPTER 6: DEPLOYMENT AND MAINTENANCE', 1)
ah('6.1 Deployment Environment', 2)
ap('Smart Mall was developed locally and deployed to a production hosting environment:')
ap('Development Environment:', bold=True)
ab('Local Server: XAMPP v8.x (Apache, MySQL, PHP 8.0+)')
ab('Operating System: Windows 10/11')
ab('Tools: VS Code, Android Studio, Chrome DevTools')
ab('Mobile Testing: Android Emulator (API 26+) + physical device')
ap('Production Environment:', bold=True)
ab('Hosting: FreeProHost shared hosting (Apache, PHP 8.0+, MySQL)')
ab('Domain: Custom domain or FreeProHost subdomain')
ab('Mobile Distribution: Signed Android APK')

ah('6.2 Deployment Process', 2)
ap('Web Application:', bold=True)
an('Export MySQL database from local XAMPP phpMyAdmin and import into production phpMyAdmin.')
an('Upload all PHP files, assets, and configuration via cPanel File Manager or FTP.')
an('Update config.php database credentials (host, username, password, dbname) for production.')
an('Update file paths and URLs to reflect the production domain.')
an('Verify file permissions: 644 for files, 755 for directories.')
an('Test all pages to ensure correct functionality in the production environment.')
ap('Mobile Application:', bold=True)
an('Update capacitor.config.json server URL to production domain.')
an('Run npx cap copy to sync web assets with the Android project.')
an('Open Android Studio: npx cap open android.')
an('Configure signing: Build > Generate Signed Bundle/APK > Create keystore > Build.')
an('Distribute the release APK for direct installation or publish to Google Play Store.')
ap('[Insert Figure: Deployment Architecture diagram showing the deployment topology with development environment (XAMPP), production hosting (FreeProHost), Android APK distribution, and client devices (browser + mobile)]', italic=True)

ah('6.3 Maintenance Plan', 2)
atab(['Frequency', 'Task', 'Responsibility'],
    [['Daily', 'Monitor server error logs', 'Admin'],
     ['Weekly', 'Review security advisories', 'Admin'],
     ['Weekly', 'Monitor disk space', 'Admin'],
     ['Monthly', 'Update PHP version', 'Developer'],
     ['Monthly', 'Review and optimize queries', 'Developer'],
     ['Monthly', 'Update Chapa API integration', 'Developer'],
     ['Quarterly', 'Full security audit', 'Developer'],
     ['Quarterly', 'Database optimization', 'Developer'],
     ['Biannually', 'Code review and refactoring', 'Developer'],
     ['Annually', 'SSL certificate renewal', 'Admin']])
apb()

# ===== CH7 =====
ah('CHAPTER 7: CONCLUSION AND RECOMMENDATIONS', 1)
ah('7.1 Summary', 2)
ap('This project successfully designed, developed, and deployed Smart Mall, a full-stack e-commerce platform with a companion mobile application specifically tailored for the Ethiopian market. The platform addresses critical gaps in the local e-commerce landscape, including multi-currency support (USD and ETB), integration with the Chapa payment gateway for ETB transactions, a comprehensive admin panel, and a Capacitor-based Android mobile application.')
ap('The system was built using PHP 8.0+ and MySQL for the backend, with HTML5, CSS3, and vanilla JavaScript for the frontend. The architecture follows a three-tier pattern with clear separation of concerns. Key technical achievements include the multi-currency engine with live exchange rate caching, secure payment integration with Chapa, CSRF-protected forms, PDO prepared statements for database security, bcrypt password hashing, and XSS prevention through proper output escaping.')
ap('The mobile application, built with Capacitor v8 and Android Studio, wraps the web platform into a native Android app with immersive full-screen display, hardware back-button navigation, and deep linking for Chapa payment callbacks. This cross-platform approach allows the mobile app to share the same codebase and business logic as the web platform while providing an enhanced mobile experience.')
ap('Testing was conducted across functional, payment, mobile, and security dimensions, with all test cases passing. The platform was deployed on FreeProHost shared hosting for production access, and the mobile app was packaged as a standalone Android APK.')

ah('7.2 Achievements', 2)
ab('Fully functional e-commerce web platform with product catalog, cart, checkout, and order management.')
ab('Multi-currency support with live exchange rate integration for USD and ETB display.')
ab('Seamless Chapa payment gateway integration enabling ETB transactions via multiple payment methods.')
ab('Companion Android mobile app with immersive display, back-button navigation, and deep linking.')
ab('Comprehensive security implementation: CSRF, prepared statements, bcrypt, XSS prevention.')
ab('Intuitive admin dashboard for products, categories, orders, and store statistics.')
ab('Responsive design with dark/light theme support for enhanced user experience.')
ab('Successful production deployment and APK distribution ready.')

ah('7.3 Limitations', 2)
ab('Single vendor only; multi-vendor marketplace functionality not implemented.')
ab('No AI-powered recommendations, personalized suggestions, or advanced analytics.')
ab('Real-time customer support features such as live chat were beyond scope.')
ab('Social media login (OAuth) not implemented; email/password registration only.')
ab('Mobile app available for Android only; iOS not targeted.')
ab('Currently English-only without multi-language localization support.')
ab('Shared hosting environment limits scalability for high traffic volumes.')

ah('7.4 Recommendations for Future Work', 2)
ab('Multi-Vendor Marketplace: Extend to support multiple vendors with individual catalogs and commission tracking.')
ab('Machine Learning: Implement product recommendations based on browsing history and purchase patterns.')
ab('Real-Time Support: Integrate live chat using WebSockets or third-party messaging APIs.')
ab('Social Login: Add OAuth authentication with Google, Facebook, and other providers.')
ab('iOS App: Develop iOS version using Capacitor\'s iOS support for the growing iPhone market.')
ab('Multi-Language: Implement localization for Amharic and other Ethiopian languages.')
ab('Cloud Migration: Migrate to AWS, Google Cloud, or VPS with auto-scaling and load balancing.')
ab('CI/CD Pipeline: Set up automated testing with PHPUnit and Playwright for code quality.')
apb()

# ===== REFERENCES =====
ah('REFERENCES', 1)
refs = [
    'P. Kotler and K. L. Keller, Marketing Management, 15th ed. Pearson, 2016.',
    'E. Turban et al., Electronic Commerce 2018: A Managerial and Social Networks Perspective, 9th ed. Springer, 2018.',
    'L. Welling and L. Thomson, PHP and MySQL Web Development, 5th ed. Addison-Wesley, 2016.',
    'M. Zandstra, PHP 8 Objects, Patterns, and Practice, 6th ed. Apress, 2021.',
    'R. Nixon, Learning PHP, MySQL & JavaScript, 6th ed. O\'Reilly Media, 2021.',
    'Chapa, "Chapa API Documentation." Available: https://developer.chapa.co/',
    'ExchangeRate-API, "API Documentation." Available: https://www.exchangerate-api.com/docs/',
    'Ionic Team, "Capacitor Documentation," 2024. Available: https://capacitorjs.com/docs/',
    'OWASP, "CSRF Prevention Cheat Sheet." Available: https://cheatsheetseries.owasp.org/',
    'OWASP, "SQL Injection Prevention Cheat Sheet." Available: https://cheatsheetseries.owasp.org/',
    'OWASP, "XSS Prevention Cheat Sheet." Available: https://cheatsheetseries.owasp.org/',
    'PHP Documentation, "PDO Manual." Available: https://www.php.net/manual/en/book.pdo.php',
    'MySQL 8.0 Reference Manual. Available: https://dev.mysql.com/doc/refman/8.0/en/',
    'Apache HTTP Server Documentation v2.4. Available: https://httpd.apache.org/docs/2.4/',
    'PHP Manual, "Password Hashing." Available: https://www.php.net/manual/en/function.password-hash.php',
]
for i, ref in enumerate(refs):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    r = p.add_run(f'[{i+1}] {ref}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
apb()

# ===== APPENDICES =====
ah('APPENDIX A: USER MANUAL', 1)
ap('Web Application:', bold=True)
ab('Browse products on the homepage; use category cards to filter; use search bar for autocomplete.')
ab('Click any product to view details, images, description, and dual-currency pricing.')
ab('Use the currency toggle (USD/ETB) to switch display; use theme toggle for dark/light mode.')
ab('Click "Add to Cart" on product pages; view cart via the cart icon; adjust quantities as needed.')
ab('Proceed to checkout: fill shipping info, select payment method (Chapa/Bank Transfer/COD), place order.')
ab('For Chapa Pay, complete payment on Chapa\'s secure page; you will return to the order confirmation.')
ap('Account Management:', bold=True)
ab('Register an account with name, email, phone, and password.')
ab('Log in to access your profile, order history, and saved information.')
ab('Use "Forgot Password" if needed; check your email for the reset link.')
ab('View and track order status in the Profile page.')
ap('Admin Panel:', bold=True)
ab('Log in with admin credentials; Dashboard shows store statistics.')
ab('Manage Products: add, edit, delete products with images and video URLs.')
ab('Manage Categories: organize products with slide images.')
ab('Manage Orders: view all orders and update their status as orders are processed.')
ap('Mobile App:', bold=True)
ab('Install the Smart Mall APK on Android 8.0+ devices.')
ab('The app provides immersive full-screen browsing of the entire platform.')
ab('Use the hardware back button for navigation; press back on home to exit.')
ab('All web features are available; Chapa payments redirect to the payment page and return on completion.')

ah('APPENDIX B: INSTALLATION GUIDE', 1)
ap('Local Development:', bold=True)
ab('Install XAMPP v8.x from https://www.apachefriends.org/')
ab('Copy project files to C:\\xampp\\htdocs\\smartmall\\')
ab('Start Apache and MySQL from XAMPP Control Panel.')
ab('Open phpMyAdmin, create database "smart_mall", import the SQL schema.')
ab('Update config.php if needed; access at http://localhost/smartmall/')
ap('Production Deployment:', bold=True)
ab('Export local MySQL database via phpMyAdmin > Export > all tables > Go.')
ab('In production cPanel phpMyAdmin: create database and user, import SQL.')
ab('Upload all files to public_html via FTP or File Manager.')
ab('Update config.php with production database credentials.')
ab('Set file permissions: 644 for files, 755 for directories.')
ap('Mobile App Build:', bold=True)
ab('Ensure Node.js/npm installed; install Capacitor CLI: npm install -g @capacitor/cli')
ab('npx cap init SmartMall com.smartmall.app && npx cap add android')
ab('Update capacitor.config.json server URL to production.')
ab('npx cap copy && npx cap open android')
ab('In Android Studio: Build > Generate Signed Bundle/APK > APK > Create keystore > Build.')

ah('APPENDIX C: SCREENSHOTS', 1)
ap('The following screenshots document key pages of the Smart Mall platform:')
add_img('01-homepage.png'); add_img('02-login.png'); add_img('03-register.png')
add_img('04-about.png'); add_img('05-contact.png'); add_img('06-category.png')
add_img('07-product-detail.png'); add_img('08-cart.png'); add_img('09-checkout.png')
add_img('10-order-confirmation.png'); add_img('11-orders.png')
add_img('12-admin-dashboard.png'); add_img('13-admin-orders.png')
add_img('14-admin-categories.png'); add_img('15-admin-add-product.png')

ah('APPENDIX D: SOURCE CODE LISTING', 1)
ap('Key files in the Smart Mall project:')
ab('config.php - Application bootstrap with session, security, and dependency loading.')
ab('db.php - PDO database connection singleton with prepared statement utilities.')
ab('currency.php - Multi-currency engine with live rate fetching and file caching.')
ab('index.php - Homepage with product grid, category filtering via AJAX.')
ab('product-detail.php - Product detail view with image gallery and video.')
ab('cart.php - Server-side shopping cart with quantity management.')
ab('checkout.php and checkout_ajax.php - Checkout form and AJAX order processing.')
ab('chapa_callback.php - Payment callback handler with verification and stock restoration.')
ab('admin/*.php - Admin dashboard, product CRUD, category management, order management.')
ab('capacitor.config.json - Mobile app configuration for Capacitor v8.')
ab('style.css - Stylesheet with CSS custom properties for theming.')
ab('script.js - Client-side JavaScript with AJAX, theme toggling, validation.')

ah('APPENDIX E: TEST CASE DETAILS', 1)
ap('Test Summary:', bold=True)
ab('Functional Tests (24/24 passed): All user flows including registration, login, product browsing, cart management, checkout with three payment methods, order history, password reset, currency/theme toggles, and admin operations.')
ab('Mobile Tests (12/12 passed): App launch, full-screen display, back-button navigation, deep linking, splash screen, cleartext restriction, Chapa payment flow, session persistence, product browsing, cart state, orientation changes.')
ab('Payment Tests (8/8 passed): Chapa initialization, success/failure callbacks, test mode, bank transfer, COD, stock restoration, multi-method switching.')
ab('Security Tests (14/14 passed): CSRF (missing/invalid token), SQL injection (login bypass/table deletion), XSS (script/event handler injection), password policy (length/complexity), bcrypt hash verification, session fixation prevention, admin access control, file upload validation (type/size).')
ap('All test results confirm that Smart Mall meets its functional and non-functional requirements.')

# ===== SAVE =====
outpath = os.path.join(os.path.dirname(__file__), 'SmartMall_Final_Year_Project.docx')
doc.save(outpath)
sz = os.path.getsize(outpath)
print(f'Document saved: {outpath}')
print(f'Size: {sz/1024:.0f} KB')
